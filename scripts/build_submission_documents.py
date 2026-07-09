from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
TABLES = ROOT / "results" / "paper_tables"
OUT = ROOT / "submission_package" / "scientific_reports_20260630_submission_docs"


PRESET = {
    "page_margin_in": 1.0,
    "body_font": "Calibri",
    "body_size_pt": 11,
    "body_after_pt": 8,
    "body_line_spacing": 1.15,
    "h1_size_pt": 16,
    "h2_size_pt": 13,
    "h3_size_pt": 12,
    "heading_blue": "2E74B5",
    "heading_dark": "1F4D78",
    "table_header_fill": "F4F6F9",
}


TABLE_INSERTIONS = {
    "CICIoT2023 fine-grained detection": "table1",
    "Rare attack classes": "table2",
    "IoTID20 independent holdout replication": "table3",
    "Leave-one-device-out N-BaIoT validation": "table4",
    "Feature contribution analysis": "table5",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if short_cell(text) else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = PRESET["body_font"]
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def short_cell(text: str) -> bool:
    return len(str(text)) <= 12 and " " not in str(text)


def fmt(value) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        return f"{value:.4f}"
    return str(value)


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    return text


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(PRESET["page_margin_in"])
    section.bottom_margin = Inches(PRESET["page_margin_in"])
    section.left_margin = Inches(PRESET["page_margin_in"])
    section.right_margin = Inches(PRESET["page_margin_in"])
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = PRESET["body_font"]
    normal.font.size = Pt(PRESET["body_size_pt"])
    normal.paragraph_format.space_after = Pt(PRESET["body_after_pt"])
    normal.paragraph_format.line_spacing = PRESET["body_line_spacing"]

    for name, size, color, before, after in [
        ("Heading 1", PRESET["h1_size_pt"], PRESET["heading_blue"], 16, 8),
        ("Heading 2", PRESET["h2_size_pt"], PRESET["heading_blue"], 12, 6),
        ("Heading 3", PRESET["h3_size_pt"], PRESET["heading_dark"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = PRESET["body_font"]
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(caption)
    run.bold = True
    run.font.name = PRESET["body_font"]
    run.font.size = Pt(10)


def set_table_geometry(table, widths_in: list[float]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def widths_for_table(table_key: str, n_cols: int) -> list[float]:
    widths = {
        "table1": [2.05, 0.85, 1.05, 0.85, 0.90, 0.80],
        "table2": [1.45, 1.30, 1.35, 1.25, 1.15],
        "table3": [2.05, 0.85, 1.05, 0.85, 0.90, 0.80],
        "table4": [0.75, 1.20, 1.15, 1.25, 1.15, 1.00],
        "table5": [0.55, 2.45, 1.50, 2.00],
    }.get(table_key)
    if widths and len(widths) == n_cols:
        return widths
    return [6.5 / n_cols] * n_cols


def add_table(doc: Document, title: str, frame: pd.DataFrame, table_key: str) -> None:
    add_caption(doc, title)
    table = doc.add_table(rows=1, cols=len(frame.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths = widths_for_table(table_key, len(frame.columns))
    for j, col in enumerate(frame.columns):
        set_cell_text(table.rows[0].cells[j], str(col), bold=True, font_size=8)
        set_cell_shading(table.rows[0].cells[j], PRESET["table_header_fill"])
    set_table_geometry(table, widths)
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(frame.columns):
            set_cell_text(cells[j], fmt(row[col]), font_size=7.5)
        set_table_geometry(table, widths)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    note.add_run("Note. Values are rounded to four decimals where applicable. Full precision is available in the accompanying CSV tables.").italic = True


def main_tables() -> dict[str, tuple[str, pd.DataFrame]]:
    t1 = pd.read_csv(TABLES / "table1_ciciot_main_results.csv")
    t1 = t1[["model", "accuracy", "balanced_accuracy", "macro_f1", "minority_recall", "pr_auc_macro"]]
    t1.columns = ["Model", "Accuracy", "Balanced accuracy", "Macro-F1", "Minority recall", "Macro PR-AUC"]

    t2 = pd.read_csv(TABLES / "table2_ciciot_rare_class_recall.csv")
    pivot = t2.pivot_table(index="class_label", columns="model", values="recall", aggfunc="first").reset_index()
    support = t2.groupby("class_label")["support"].max().reset_index()
    t2 = pivot.merge(support, on="class_label", how="left")
    cols = [c for c in ["class_label", "Flat LightGBM", "Weighted XGBoost", "Ensemble", "support"] if c in t2.columns]
    t2 = t2[cols]
    t2.columns = ["Class", "Flat LightGBM recall", "Weighted XGBoost recall", "Ensemble recall", "Support"]

    t3 = pd.read_csv(TABLES / "table3_iotid20_replication.csv")
    t3 = t3[["model", "accuracy", "balanced_accuracy", "macro_f1", "minority_recall", "pr_auc_macro"]]
    t3.columns = ["Model", "Accuracy", "Balanced accuracy", "Macro-F1", "Minority recall", "Macro PR-AUC"]

    t4 = pd.read_csv(TABLES / "table4_nbaiot_unseen_device_summary.csv")
    t4 = t4[["task", "mean_macro_f1", "min_macro_f1", "mean_present_macro_f1", "min_present_class_recall", "mean_balanced_accuracy"]]
    t4.columns = ["Task", "Conservative Macro-F1 mean", "Conservative Macro-F1 min", "Present-class Macro-F1 mean", "Minimum present-class recall", "Balanced accuracy mean"]

    t5 = pd.read_csv(TABLES / "table5_ciciot_shap_top_features.csv").head(10)
    t5 = t5[["rank", "feature", "mean_abs_shap", "lightgbm_gain_importance"]]
    t5.columns = ["Rank", "Feature", "Mean absolute SHAP", "LightGBM gain"]

    return {
        "table1": ("Table 1. CICIoT2023 main fine-grained detection results.", t1),
        "table2": ("Table 2. Rare CICIoT2023 class recall.", t2),
        "table3": ("Table 3. IoTID20 independent holdout replication.", t3),
        "table4": ("Table 4. N-BaIoT leave-one-device-out validation summary.", t4),
        "table5": ("Table 5. CICIoT2023 LightGBM component SHAP feature contributions.", t5),
    }


def add_paragraph_with_basic_bold(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    text = clean_inline(text)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


def build_from_markdown(md_path: Path, out_path: Path, include_tables: bool = False) -> None:
    doc = Document()
    configure_doc(doc)
    tables = main_tables() if include_tables else {}
    pending_table: str | None = None
    added_tables: set[str] = set()
    in_code = False

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            continue
        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(clean_inline(line[2:]))
            run.bold = True
            run.font.name = PRESET["body_font"]
            run.font.size = Pt(18)
            continue
        if line.startswith("## "):
            doc.add_heading(clean_inline(line[3:]), level=1)
            continue
        if line.startswith("### "):
            heading = clean_inline(line[4:])
            doc.add_heading(heading, level=2)
            pending_table = TABLE_INSERTIONS.get(heading) if include_tables else None
            continue
        if line.startswith("- "):
            add_paragraph_with_basic_bold(doc, line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            add_paragraph_with_basic_bold(doc, re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            add_paragraph_with_basic_bold(doc, line)

        if pending_table and pending_table not in added_tables:
            title, frame = tables[pending_table]
            add_table(doc, title, frame, pending_table)
            added_tables.add(pending_table)
            pending_table = None

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def build_cover_letter() -> None:
    build_from_markdown(MANUSCRIPT / "cover_letter_scientific_reports.md", OUT / "cover_letter_scientific_reports.docx")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    build_from_markdown(MANUSCRIPT / "scientific_reports_draft.md", OUT / "scientific_reports_main_text.docx", include_tables=True)
    build_from_markdown(MANUSCRIPT / "supplementary_materials.md", OUT / "scientific_reports_supplementary_information.docx")
    build_cover_letter()
    print(f"Wrote DOCX files to {OUT}")


if __name__ == "__main__":
    main()
